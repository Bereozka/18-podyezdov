from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from typing import Union
from typing_extensions import Literal


def create_word_file(
    work_data,
    material_data,
    total_price,
    template_path,
    final_file_path,
):
    doc = Document(template_path)

    table = doc.tables[0]
    insert_info_into_table(work_data, table)
    insert_info_into_table(material_data, doc.tables[1])
    add_total_price(doc.paragraphs, total_price)

    doc.save(final_file_path)


def add_total_price(
    paragraphs,
    total_price,
):
    for paragraph in paragraphs:
        if "ВСЕГО" in paragraph.text:
            paragraph.text = ""
            insert_formatted_text(
                paragraph,
                total_price,
                alignment="center",
                bold=True,
                font_size=12,
            )


def insert_info_into_table(
    data,
    table,
):
    for item_index, item in enumerate(data):
        row = table.add_row()
        index = str(item_index + 1)
        if item["type"] == "filled":
            insert_filled_row(
                row,
                index,
                item["data"]["name"],
                item["data"]["units"],
                item["data"]["count"],
                item["data"]["price"],
                item["data"]["total"],
            )
        elif item["type"] == "subtitle":
            insert_subtitle_row(
                row,
                index,
                item["data"]["title"],
            )
        elif item["type"] == "total":
            insert_total_row(
                row,
                index,
                item["data"]["value"],
            )


def insert_filled_row(
    row,
    index,
    name,
    units,
    count,
    price,
    total,
):
    insert_formatted_text(
        row.cells[0].paragraphs[0],
        index,
        alignment="center",
        bold=True,
    )
    insert_formatted_text(
        row.cells[1].paragraphs[0],
        name,
    )

    items = [units, count, price, total]
    for i in range(2, 6):
        insert_formatted_text(
            row.cells[i].paragraphs[0],
            items[i - 2],
            alignment="center",
        )


def insert_subtitle_row(
    row,
    index,
    text,
):
    insert_formatted_text(
        row.cells[0].paragraphs[0],
        index,
        alignment="center",
        bold=True,
    )

    row.cells[1].merge(row.cells[-1])
    insert_formatted_text(
        row.cells[1].paragraphs[0],
        text,
        alignment="center",
        bold=True,
    )


def insert_total_row(
    row,
    index,
    text,
):
    insert_formatted_text(
        row.cells[0].paragraphs[0],
        index,
        alignment="center",
        bold=True,
    )
    row.cells[1].merge(row.cells[-1])
    insert_formatted_text(
        row.cells[1].paragraphs[0],
        text,
        alignment="center",
        bold=True,
    )


def insert_formatted_text(
    paragraph: Paragraph,
    value: str,
    bold: bool = False,
    alignment: Literal["left", "center", "right"] = "left",
    font_size: Union[int, float] = 10.5,
) -> None:
    """Formats the paragraph text

    Parameters
    ----------
    paragraph : Paragraph
        An object of the Paragraph class
    value : str
        The text to be written in the paragraph
    bold : bool, default: False
        Text weight
    alignment : {"left", "center", "right"}, default: "left"
        Text alignment
    font_size : int, default: 10.5
        Text font size

    Raises
    ------
    ValueError
        If one of the parameters does not match its type

    """

    if not isinstance(paragraph, Paragraph):
        raise ValueError("Parameter paragraph should be type Paragraph")
    elif not isinstance(value, str):
        raise ValueError("Parameter value should be str")
    elif not isinstance(bold, bool):
        raise ValueError("Parameter bold should be bool")
    elif alignment not in ("left", "center", "right"):
        raise ValueError("Parameter alignment should be str")
    elif not isinstance(font_size, (int, float)):
        raise ValueError("Parameter font_size should be int or float")

    # bold
    runner = paragraph.add_run(value)
    runner.bold = bold

    # alignment
    if alignment == "left":
        paragraph.alignment = 0
    elif alignment == "center":
        paragraph.alignment = 1
    elif alignment == "right":
        paragraph.alignment = 2

    # font size
    runner.font.size = Pt(font_size)
